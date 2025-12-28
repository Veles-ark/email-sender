import re
import time
import smtplib
import ssl
import json
import threading
import sys
# === PATCH_RANDOM_DELAY_START ===
import random
# === PATCH_RANDOM_DELAY_END ===
from pathlib import Path

from docx import Document                  # pip install python-docx
from openpyxl import Workbook, load_workbook  # pip install openpyxl
from openpyxl.styles import PatternFill

from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from email.header import Header

import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk  # для Progressbar


# =============== НАСТРОЙКИ ПО УМОЛЧАНИЮ =================

DEFAULT_SMTP_SERVER = "smtp.mail.ru"
DEFAULT_SMTP_PORT = 465  # для SMTP_SSL
# === PATCH_SMTP_SETTINGS_START ===
DEFAULT_SMTP_MODE = "SSL"  # SSL | STARTTLS | PLAIN
# === PATCH_SMTP_SETTINGS_END ===

DEFAULT_FROM_EMAIL = "example@mail.ru"  # можно оставить пустым ""

# Определяем папку, где лежит программа (и .py, и .exe)
if getattr(sys, "frozen", False):
    # Запущено как упакованный EXE
    APP_DIR = Path(sys.argv[0]).resolve().parent
else:
    # Запущено как обычный .py
    APP_DIR = Path(__file__).resolve().parent

# Файл настроек (сохраняем почту, пароль, тему и т.п.) – рядом с программой
SETTINGS_FILE = APP_DIR / "email_tool_settings.json"

# === PATCH_FILE_LOG_START ===
LOG_FILE = APP_DIR / "email_tool.log"
# === PATCH_FILE_LOG_END ===

# Файл с текстом сообщения по умолчанию, если поле в GUI пустое – тоже рядом с программой
MESSAGE_DOCX_NAME = "сообщение.docx"


# ======================================================================


# ======== ЧТЕНИЕ АДРЕСОВ ИЗ WORD И ЗАПИСЬ В EXCEL =========

def extract_emails_from_docx(docx_path: Path):
    """Возвращает список всех e-mail адресов, найденных в документе."""
    doc = Document(docx_path)
    emails = []

    email_pattern = re.compile(
        r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+(?=[\s,;:\)\]\}?!]|$)'
    )

    for para in doc.paragraphs:
        found = email_pattern.findall(para.text)
        emails.extend(
            [email.replace("\n", "").replace("\r", "").strip().rstrip(".")
             for email in found]
        )

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                found = email_pattern.findall(cell.text)
                emails.extend(
                    [email.replace("\n", "").replace("\r", "").strip().rstrip(".")
                     for email in found]
                )

    return emails


def save_emails_to_excel(emails, excel_path: Path):
    """Сохраняет список e-mail адресов в Excel, один адрес в строку (столбец A)."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Emails"

    row = 1
    for email in emails:
        ws.cell(row=row, column=1, value=email)
        row += 1

    wb.save(excel_path)


# ======== ЗАГРУЗКА АДРЕСОВ ИЗ EXCEL ДЛЯ РАССЫЛКИ =========

def load_emails_with_rows(path: Path):
    """
    Открывает Excel, возвращает:
    - workbook
    - active sheet
    - список (row_index, email) по первой колонке.
    """
    if not path.exists():
        raise FileNotFoundError(f"Файл с адресами не найден: {path}")

    wb = load_workbook(path)
    ws = wb.active

    entries = []
    for row_idx, row in enumerate(ws.iter_rows(min_row=1, max_col=1), start=1):
        cell_value = row[0].value
        if cell_value is None:
            continue
        email = str(cell_value).strip()
        if email:
            entries.append((row_idx, email))

    return wb, ws, entries


# ============= ТЕКСТ СООБЩЕНИЯ ИЗ DOCX (если поле пустое) =============

def load_message_from_docx():
    """Читает текст из MESSAGE_DOCX_NAME (сообщение.docx) рядом с программой (и .py, и .exe)."""
    path = APP_DIR / MESSAGE_DOCX_NAME
    if not path.exists():
        raise FileNotFoundError(f"Файл с текстом сообщения не найден: {path}")

    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]
    text = "\n".join(paragraphs).strip()
    return text



# ================ СОЗДАНИЕ ПИСЬМА =================

def create_message(from_email: str,
                   to_email: str,
                   subject: str,
                   body: str,
                   attachment_paths: list[Path] | None) -> MIMEMultipart:
    """Создаёт MIME-сообщение с текстом и несколькими вложениями."""
    msg = MIMEMultipart()
    msg["From"] = from_email
    msg["To"] = to_email
    msg["Subject"] = Header(subject, "utf-8")

    msg.attach(MIMEText(body, "plain", "utf-8"))

    if attachment_paths:
        for attachment_path in attachment_paths:
            if attachment_path.exists():
                with attachment_path.open("rb") as f:
                    part = MIMEBase("application", "octet-stream")
                    part.set_payload(f.read())
                encoders.encode_base64(part)
                # === PATCH_RUS_FILENAME_START ===
                # part.add_header(
                #     "Content-Disposition",
                #     f'attachment; filename="{attachment_path.name}"'
                # )
                part.add_header(
                    "Content-Disposition",
                    "attachment",
                    filename=("utf-8", "", attachment_path.name)
                )
                # === PATCH_RUS_FILENAME_END ===
                msg.attach(part)
            else:
                log(f"ВНИМАНИЕ: файл вложения не найден: {attachment_path}. Этот файл не будет прикреплён.")
    return msg



# ===================== ЛОГИРОВАНИЕ В GUI =======================

def log(message: str):
    # === PATCH_FILE_LOG_START ===
    try:
        with LOG_FILE.open("a", encoding="utf-8") as f:
            f.write(message + "\n")
    except Exception:
        pass
    # === PATCH_FILE_LOG_END ===

    log_text.configure(state="normal")
    log_text.insert("end", message + "\n")
    log_text.see("end")
    log_text.configure(state="disabled")
    root.update_idletasks()


def thread_log(message: str):
    """Безопасное логирование из фонового потока."""
    root.after(0, lambda m=message: log(m))


# ===================== ПРОГРЕСС-БАР =======================

def init_progress(total: int):
    """Инициализация прогресс-бара (вызывается через root.after)."""
    def _init():
        progress_bar["maximum"] = max(total, 1)
        progress_bar["value"] = 0
    root.after(0, _init)


def update_progress(current: int, total: int):
    """Обновление прогресса (вызывается из фонового потока)."""
    def _update():
        progress_bar["maximum"] = max(total, 1)
        progress_bar["value"] = current
    root.after(0, _update)


def reset_progress():
    """Сброс прогресса (в конец или в 0 — на выбор)."""
    def _reset():
        progress_bar["value"] = 0
    root.after(0, _reset)


# ===================== РАБОТА С НАСТРОЙКАМИ =======================

def save_settings():
    """Сохраняет текущие значения полей в SETTINGS_FILE."""
    try:
        data = {
            "docx_path": docx_var.get(),
            "excel_save_path": excel_save_var.get(),
            "excel_send_path": excel_send_var.get(),
            # === PATCH_SMTP_SETTINGS_START ===
            "smtp_server": smtp_server_var.get(),
            "smtp_port": smtp_port_var.get(),
            "smtp_mode": smtp_mode_var.get(),
            "smtp_user": smtp_user_var.get(),
            # === PATCH_SMTP_SETTINGS_END ===
            "from_email": from_email_var.get(),
            "password": password_var.get(),
            "subject": subject_var.get(),
            "body": body_text.get("1.0", "end"),
            "attach_path": attach_var.get(),
            # === PATCH_RANDOM_DELAY_START ===
            "delay_from_minutes": delay_from_var.get(),
            "delay_to_minutes": delay_to_var.get(),
            # === PATCH_RANDOM_DELAY_END ===
            "delay_minutes": delay_var.get(),
        }

        SETTINGS_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        log(f"Настройки сохранены в {SETTINGS_FILE}")
    except Exception as e:
        log(f"Ошибка сохранения настроек: {e}")


def load_settings():
    """Загружает значения полей из SETTINGS_FILE, если он существует."""
    if not SETTINGS_FILE.exists():
        return
    try:
        data = json.loads(SETTINGS_FILE.read_text(encoding="utf-8"))

        docx_var.set(data.get("docx_path", ""))
        excel_save_var.set(data.get("excel_save_path", ""))
        excel_send_var.set(data.get("excel_send_path", ""))
        # === PATCH_SMTP_SETTINGS_START ===
        smtp_server_var.set(data.get("smtp_server", DEFAULT_SMTP_SERVER))
        smtp_port_var.set(data.get("smtp_port", str(DEFAULT_SMTP_PORT)))
        smtp_mode_var.set(data.get("smtp_mode", DEFAULT_SMTP_MODE))
        smtp_user_var.set(data.get("smtp_user", ""))
        # === PATCH_SMTP_SETTINGS_END ===

        from_email_var.set(data.get("from_email", DEFAULT_FROM_EMAIL))
        password_var.set(data.get("password", ""))

        subject_var.set(data.get("subject", ""))
        attach_var.set(data.get("attach_path", ""))
        delay_var.set(data.get("delay_minutes", "5"))

        # === PATCH_RANDOM_DELAY_START ===
        legacy_delay = data.get("delay_minutes", "5")
        delay_from_var.set(data.get("delay_from_minutes", str(legacy_delay)))
        delay_to_var.set(data.get("delay_to_minutes", str(legacy_delay)))
        # === PATCH_RANDOM_DELAY_END ===

        body_text.delete("1.0", "end")
        body_text.insert("1.0", data.get("body", ""))

        log(f"Настройки загружены из {SETTINGS_FILE}")
    except Exception as e:
        log(f"Ошибка загрузки настроек: {e}")


# ===================== GUI ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ =======================

def gui_select_docx_file():
    file_path = filedialog.askopenfilename(
        title="Выберите Word-файл (.docx)",
        filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
    )
    if not file_path:
        return

    docx_var.set(file_path)
    docx_path = Path(file_path)
    default_excel = docx_path.parent / "adres.xlsx"
    excel_save_var.set(str(default_excel))


def gui_select_excel_save_file():
    file_path = filedialog.asksaveasfilename(
        title="Сохранить Excel-файл как...",
        defaultextension=".xlsx",
        filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")]
    )
    if not file_path:
        return
    excel_save_var.set(file_path)


def gui_select_excel_send_file():
    file_path = filedialog.askopenfilename(
        title="Выберите Excel-файл для рассылки",
        filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")]
    )
    if not file_path:
        return
    excel_send_var.set(file_path)


def gui_select_attach_file():
    """Выбор одного или нескольких файлов-вложений."""
    file_paths = filedialog.askopenfilenames(
        title="Выберите файл(ы) для вложения",
        filetypes=[("All files", "*.*")]
    )
    if not file_paths:
        return

    # сохраняем все пути в одну строку, разделяя ';'
    attach_var.set("; ".join(file_paths))



def gui_run_extraction():
    docx_path_str = docx_var.get().strip()
    excel_path_str = excel_save_var.get().strip()

    if not docx_path_str:
        messagebox.showerror("Ошибка", "Не выбран Word-файл (.docx).")
        return

    docx_path = Path(docx_path_str)
    if not docx_path.exists():
        messagebox.showerror("Ошибка", f"Файл не найден:\n{docx_path}")
        return

    if docx_path.suffix.lower() != ".docx":
        messagebox.showerror("Ошибка", "Нужно выбрать файл формата .docx.")
        return

    if not excel_path_str:
        messagebox.showerror("Ошибка", "Не указан путь для Excel-файла.")
        return

    excel_path = Path(excel_path_str)

    try:
        emails = extract_emails_from_docx(docx_path)
        if not emails:
            messagebox.showinfo("Результат", "E-mail адреса не найдены.")
            return

        save_emails_to_excel(emails, excel_path)
        messagebox.showinfo(
            "Готово",
            f"Найдено адресов: {len(emails)}\n"
            f"Результат сохранён в:\n{excel_path}"
        )
        log(f"Импортировано адресов: {len(emails)} в файл {excel_path}")
    except Exception as e:
        messagebox.showerror("Ошибка выполнения", str(e))
        log(f"Ошибка импорта адресов: {e}")


def gui_run_sending():
    """Собираем данные из GUI и запускаем фоновый поток отправки."""
    excel_path_str = excel_send_var.get().strip()
    subject = subject_var.get().strip()
    body = body_text.get("1.0", "end").strip()
    attach_paths_str = attach_var.get().strip()
    delay_minutes_str = delay_var.get().strip()

    # === PATCH_RANDOM_DELAY_START ===
    delay_from_minutes_str = delay_from_var.get().strip()
    delay_to_minutes_str = delay_to_var.get().strip()
    # === PATCH_RANDOM_DELAY_END ===

    # === PATCH_SMTP_SETTINGS_START ===
    smtp_server = smtp_server_var.get().strip()
    smtp_port_str = smtp_port_var.get().strip()
    smtp_mode = smtp_mode_var.get().strip()
    smtp_user = smtp_user_var.get().strip()
    # === PATCH_SMTP_SETTINGS_END ===

    from_email = from_email_var.get().strip()
    password = password_var.get().strip()

    # === PATCH_SMTP_SETTINGS_START ===
    if not smtp_server:
        messagebox.showerror("Ошибка", "Не указан SMTP сервер.")
        return

    try:
        smtp_port = int(smtp_port_str)
        if smtp_port <= 0 or smtp_port > 65535:
            raise ValueError
    except ValueError:
        messagebox.showerror("Ошибка", "Неверный SMTP порт.")
        return

    if smtp_mode not in ("SSL", "STARTTLS", "PLAIN"):
        messagebox.showerror("Ошибка", "Неверный режим SMTP (SSL / STARTTLS / PLAIN).")
        return

    if not smtp_user:
        smtp_user = from_email
    # === PATCH_SMTP_SETTINGS_END ===


    if not from_email:
        messagebox.showerror("Ошибка", "Не указан адрес отправителя.")
        return

    if not password:
        messagebox.showerror("Ошибка", "Не указан пароль для почты отправителя.")
        return


    if not excel_path_str:
        messagebox.showerror("Ошибка", "Не указан Excel-файл для рассылки.")
        return

    if not subject:
        messagebox.showerror("Ошибка", "Не указана тема письма.")
        return

    # Если текст пустой — берём из сообщение.docx
    if not body:
        try:
            body = load_message_from_docx()
            log(f"Поле 'ТЕКСТ СООБЩЕНИЯ' пустое. Текст загружен из {MESSAGE_DOCX_NAME}.")
        except Exception as e:
            messagebox.showerror(
                "Ошибка",
                f"Поле 'ТЕКСТ СООБЩЕНИЯ' пустое и не удалось загрузить текст из {MESSAGE_DOCX_NAME}:\n{e}"
            )
            log(f"Ошибка загрузки текста из {MESSAGE_DOCX_NAME}: {e}")
            return

    try:
        delay_minutes = float(delay_minutes_str) if delay_minutes_str else 5.0
        if delay_minutes < 0:
            raise ValueError
    except ValueError:
        messagebox.showerror("Ошибка", "Неверное значение интервала в минутах.")
        return

    # === PATCH_RANDOM_DELAY_START ===
    try:
        delay_from_minutes = float(delay_from_minutes_str) if delay_from_minutes_str else 1.0
        delay_to_minutes = float(delay_to_minutes_str) if delay_to_minutes_str else 5.0
        if delay_from_minutes < 0 or delay_to_minutes < 0:
            raise ValueError
        if delay_to_minutes < delay_from_minutes:
            messagebox.showerror("Ошибка", "Поле 'до' должно быть больше или равно полю 'от'.")
            return
    except ValueError:
        messagebox.showerror("Ошибка", "Неверные значения интервала (от/до) в минутах.")
        return
    # === PATCH_RANDOM_DELAY_END ===


    save_settings()
    send_button.config(state="disabled")
    reset_progress()

    t = threading.Thread(
        target=send_worker,
        # === PATCH_RANDOM_DELAY_START ===
        args=(excel_path_str, subject, body, attach_paths_str, delay_from_minutes, delay_to_minutes, smtp_server, smtp_port, smtp_mode, smtp_user, from_email, password),
        # === PATCH_RANDOM_DELAY_END ===
        daemon=True
    )
    t.start()


def send_worker(excel_path_str: str,
                subject: str,
                body: str,
                attach_paths_str: str,
                # === PATCH_RANDOM_DELAY_START ===
                delay_from_minutes: float,
                delay_to_minutes: float,
                # === PATCH_RANDOM_DELAY_END ===
                # delay_minutes: float,
                # === PATCH_SMTP_SETTINGS_START ===
                smtp_server: str,
                smtp_port: int,
                smtp_mode: str,
                smtp_user: str,
                # === PATCH_SMTP_SETTINGS_END ===
                from_email: str,
                password: str):
    """Фоновый поток: выполняет рассылку, не блокируя GUI."""
    # === PATCH_RANDOM_DELAY_START ===
    delay_from_seconds = int(delay_from_minutes * 60)
    delay_to_seconds = int(delay_to_minutes * 60)
    # === PATCH_RANDOM_DELAY_END ===
    # delay_seconds = int(delay_minutes * 60)
    excel_path = Path(excel_path_str)

    # Преобразуем строку с путями в список Path
    if attach_paths_str:
        attachment_paths = [
            Path(p.strip()) for p in attach_paths_str.split(";") if p.strip()
        ]
    else:
        attachment_paths = []


    try:
        wb, ws, entries = load_emails_with_rows(excel_path)
    except Exception as e:
        thread_log(f"Ошибка чтения Excel: {e}")
        root.after(0, lambda: messagebox.showerror("Ошибка чтения Excel", str(e)))
        root.after(0, lambda: send_button.config(state="normal"))
        return

    if not entries:
        thread_log("В Excel-файле не найдено ни одного адреса.")
        root.after(0, lambda: messagebox.showinfo("Результат", "В Excel-файле не найдено ни одного адреса."))
        root.after(0, lambda: send_button.config(state="normal"))
        return

    total = len(entries)
    init_progress(total)

    red_fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")

    thread_log(f"Найдено адресов для отправки: {total}")
    thread_log(f"Отправитель: {from_email}")
    if attachment_paths:
        thread_log("Вложения:")
        for p in attachment_paths:
            thread_log(f"  - {p}")
    else:
        thread_log("Вложения: нет")

    # === PATCH_RANDOM_DELAY_START ===
    thread_log(f"Интервал между отправками (рандом): от {delay_from_minutes} до {delay_to_minutes} минут")
    # === PATCH_RANDOM_DELAY_END ===
    # thread_log(f"Интервал между отправками: {delay_minutes} минут")
    thread_log("Начинаем рассылку...\n")

    try:
        # === PATCH_SMTP_PER_EMAIL_START ===
        def smtp_connect_and_login():
            thread_log("SMTP: создаём SSL-контекст")
            context = ssl.create_default_context()

            thread_log(f"SMTP: сервер={smtp_server} порт={smtp_port} режим={smtp_mode}")

            if smtp_mode == "SSL":
                thread_log("SMTP: подключение через SMTP_SSL")
                s = smtplib.SMTP_SSL(smtp_server, smtp_port, context=context)

            else:
                thread_log("SMTP: подключение через SMTP (без SSL на старте)")
                s = smtplib.SMTP(smtp_server, smtp_port)

                if smtp_mode == "STARTTLS":
                    thread_log("SMTP: выполняем starttls()")
                    s.starttls(context=context)
                    thread_log("SMTP: starttls() успешно")

            thread_log(f"SMTP: выполняем login() пользователем: {smtp_user}")
            s.login(smtp_user, password)

            thread_log("SMTP: login успешен")
            return s


        # === PATCH_SMTP_PER_EMAIL_END ===

        processed = 0


        for i, (row_idx, to_email) in enumerate(entries, start=1):

            thread_log(f"[{i}/{total}] === НАЧАЛО ОБРАБОТКИ АДРЕСА === {to_email}")


            previous_status = ws.cell(row=row_idx, column=2).value
            if previous_status and str(previous_status).strip().lower() == "отправлено":
                thread_log(f"[{i}/{total}] Уже отправлено ранее — пропуск: {to_email}")
                processed += 1
                update_progress(processed, total)
                continue

            status_text = "не отправлено"
            success = False

            try:
                thread_log(f"[{i}/{total}] Создаём MIME-сообщение")
                msg = create_message(from_email, to_email, subject, body, attachment_paths)

                thread_log(f"[{i}/{total}] Открываем SMTP-соединение")
                server = smtp_connect_and_login()

                try:
                    thread_log(f"[{i}/{total}] Отправляем письмо")
                    server.sendmail(from_email, [to_email], msg.as_string())
                    thread_log(f"[{i}/{total}] SMTP: sendmail() выполнен успешно")
                finally:
                    try:
                        thread_log(f"[{i}/{total}] Закрываем SMTP-соединение")
                        server.quit()
                        thread_log(f"[{i}/{total}] SMTP-соединение закрыто")
                    except Exception as e:
                        thread_log(f"[{i}/{total}] ОШИБКА при закрытии SMTP: {e}")

                thread_log(f"[{i}/{total}] ПИСЬМО УСПЕШНО ОТПРАВЛЕНО: {to_email}")
                status_text = "отправлено"
                success = True


            except Exception as e:
                thread_log(f"[{i}/{total}] !!! ОШИБКА НА ЭТАПЕ ОТПРАВКИ !!!")
                thread_log(f"[{i}/{total}] Адрес: {to_email}")
                thread_log(f"[{i}/{total}] Текст ошибки: {e}")
                status_text = "не отправлено"
                success = False



            ws.cell(row=row_idx, column=2, value=status_text)

            if not success:
                addr_cell = ws.cell(row=row_idx, column=1)
                addr_cell.fill = red_fill

            wb.save(excel_path)

            processed += 1
            update_progress(processed, total)

            if i < total:
                # === PATCH_RANDOM_DELAY_START ===
                wait_seconds = random.randint(delay_from_seconds, max(delay_to_seconds, delay_from_seconds))
                if success:
                    thread_log(f"[{i}/{total}] Пауза {wait_seconds} секунд до следующего письма")
                else:
                    thread_log(f"[{i}/{total}] Пауза {wait_seconds} секунд после ошибки")

                time.sleep(wait_seconds)
                # === PATCH_RANDOM_DELAY_END ===


        thread_log("\n=== РАССЫЛКА ЗАВЕРШЕНА ===")
        thread_log(f"Всего адресов: {total}")
        thread_log("Результаты сохранены в Excel")

        root.after(0, lambda: messagebox.showinfo("Готово", "Рассылка завершена. Подробности смотрите в файле Excel."))


    except Exception as e:
        thread_log(f"Ошибка SMTP: {e}")
        root.after(0, lambda: messagebox.showerror("Ошибка SMTP", str(e)))
    finally:
        root.after(0, lambda: send_button.config(state="normal"))


# ===================== ЗАПУСК GUI =======================

def main():
    global root, docx_var, excel_save_var, excel_send_var
    global subject_var, attach_var, delay_var
    # === PATCH_SMTP_SETTINGS_START ===
    global smtp_server_var, smtp_port_var, smtp_mode_var, smtp_user_var
    # === PATCH_SMTP_SETTINGS_END ===

    # === PATCH_RANDOM_DELAY_START ===
    global delay_from_var, delay_to_var
    # === PATCH_RANDOM_DELAY_END ===
    global body_text, from_email_var, password_var, log_text, send_button, progress_bar

    root = tk.Tk()
    root.title("Импорт e-mail из Word и рассылка")

    # === PATCH_RESIZE_START ===
    root.grid_rowconfigure(1, weight=1)
    root.grid_columnconfigure(0, weight=1)
    # === PATCH_RESIZE_END ===


    # ---- Прогресс-бар сверху ----
    progress_bar = ttk.Progressbar(root, orient="horizontal", length=400, mode="determinate")
    progress_bar.grid(row=0, column=0, columnspan=3, padx=5, pady=5, sticky="we")

    # Основной фрейм под все остальные элементы
    frame = tk.Frame(root)
    frame.grid(row=1, column=0, columnspan=3, sticky="nsew")

    # === PATCH_RESIZE_START ===
    frame.grid_columnconfigure(1, weight=1)
    frame.grid_rowconfigure(11, weight=1)  # ТЕКСТ СООБЩЕНИЯ (Text)
    frame.grid_rowconfigure(16, weight=1)  # ПРОЦЕСС ОТПРАВКИ (Text)
    # === PATCH_RESIZE_END ===



    docx_var = tk.StringVar()
    excel_save_var = tk.StringVar()
    excel_send_var = tk.StringVar()
    subject_var = tk.StringVar()
    attach_var = tk.StringVar()
    delay_var = tk.StringVar(value="5")
    # === PATCH_RANDOM_DELAY_START ===
    delay_from_var = tk.StringVar(value="1")
    delay_to_var = tk.StringVar(value="5")
    # === PATCH_RANDOM_DELAY_END ===
    # === PATCH_SMTP_SETTINGS_START ===
    smtp_server_var = tk.StringVar(value=DEFAULT_SMTP_SERVER)
    smtp_port_var = tk.StringVar(value=str(DEFAULT_SMTP_PORT))
    smtp_mode_var = tk.StringVar(value=DEFAULT_SMTP_MODE)
    smtp_user_var = tk.StringVar(value="")
    # === PATCH_SMTP_SETTINGS_END ===
    from_email_var = tk.StringVar(value=DEFAULT_FROM_EMAIL)
    password_var = tk.StringVar()


    # ---------- Блок 1: Импорт адресов из Word ----------

    tk.Label(frame, text="Word-файл (.docx):").grid(row=0, column=0, sticky="w", padx=5, pady=5)
    tk.Entry(frame, textvariable=docx_var, width=50).grid(row=0, column=1, padx=5, pady=5)
    tk.Button(frame, text="Обзор...", command=gui_select_docx_file).grid(row=0, column=2, padx=5, pady=5)

    tk.Label(frame, text="Excel для сохранения адресов (.xlsx):").grid(row=1, column=0, sticky="w", padx=5, pady=5)
    tk.Entry(frame, textvariable=excel_save_var, width=50).grid(row=1, column=1, padx=5, pady=5)
    tk.Button(frame, text="Выбрать...", command=gui_select_excel_save_file).grid(row=1, column=2, padx=5, pady=5)

    tk.Button(frame, text="Извлечь адреса из Word в Excel", command=gui_run_extraction, width=40).grid(
        row=2, column=0, columnspan=3, pady=10
    )

    # ---------- Блок 2: Excel для рассылки ----------

    tk.Label(frame, text="Excel для рассылки (.xlsx):").grid(row=3, column=0, sticky="w", padx=5, pady=5)
    tk.Entry(frame, textvariable=excel_send_var, width=50).grid(row=3, column=1, padx=5, pady=5)
    tk.Button(frame, text="Обзор...", command=gui_select_excel_send_file).grid(row=3, column=2, padx=5, pady=5)

    # Разделитель
    tk.Label(frame, text="").grid(row=4, column=0, pady=5)

    # ---------- Блок 3: Данные отправителя ----------

    # === PATCH_SMTP_SETTINGS_START ===
    tk.Label(frame, text="SMTP сервер:").grid(row=5, column=0, sticky="w", padx=5, pady=5)
    tk.Entry(frame, textvariable=smtp_server_var, width=50).grid(
        row=5, column=1, columnspan=2, padx=5, pady=5, sticky="we"
    )

    tk.Label(frame, text="SMTP порт:").grid(row=6, column=0, sticky="w", padx=5, pady=5)
    tk.Entry(frame, textvariable=smtp_port_var, width=10).grid(row=6, column=1, sticky="w", padx=5, pady=5)

    tk.Label(frame, text="Режим:").grid(row=6, column=1, sticky="w", padx=95, pady=5)
    smtp_mode_combo = ttk.Combobox(frame, textvariable=smtp_mode_var, values=["SSL", "STARTTLS", "PLAIN"], width=10, state="readonly")
    smtp_mode_combo.grid(row=6, column=2, sticky="w", padx=5, pady=5)

    tk.Label(frame, text="SMTP логин (если нужен):").grid(row=7, column=0, sticky="w", padx=5, pady=5)
    tk.Entry(frame, textvariable=smtp_user_var, width=50).grid(
        row=7, column=1, columnspan=2, padx=5, pady=5, sticky="we"
    )
    # === PATCH_SMTP_SETTINGS_END ===

    tk.Label(frame, text="Почта отправителя:").grid(row=8, column=0, sticky="w", padx=5, pady=5)
    tk.Entry(frame, textvariable=from_email_var, width=50).grid(
        row=8, column=1, columnspan=2, padx=5, pady=5, sticky="we"
    )

    tk.Label(frame, text="Пароль:").grid(row=9, column=0, sticky="w", padx=5, pady=5)
    tk.Entry(frame, textvariable=password_var, width=50, show="*").grid(
        row=9, column=1, columnspan=2, padx=5, pady=5, sticky="we"
    )

    # ---------- Блок 4: Настройки письма и отправка ----------

    tk.Label(frame, text="ТЕМА:").grid(row=10, column=0, sticky="w", padx=5, pady=5)
    tk.Entry(frame, textvariable=subject_var, width=50).grid(
        row=10, column=1, columnspan=2, padx=5, pady=5, sticky="we"
    )

    tk.Label(frame, text="ТЕКСТ СООБЩЕНИЯ:").grid(row=11, column=0, sticky="nw", padx=5, pady=5)
    body_text = tk.Text(frame, width=60, height=8)
    # === PATCH_RESIZE_START ===
    body_text.grid(row=11, column=1, columnspan=2, padx=5, pady=5, sticky="nsew")
    # === PATCH_RESIZE_END ===


    tk.Label(frame, text="Вложение:").grid(row=12, column=0, sticky="w", padx=5, pady=5)
    tk.Entry(frame, textvariable=attach_var, width=50).grid(row=12, column=1, padx=5, pady=5)
    tk.Button(frame, text="Прикрепить.", command=gui_select_attach_file).grid(row=12, column=2, padx=5, pady=5)

    # === PATCH_RANDOM_DELAY_START ===
    tk.Label(frame, text="Минут между отправками (от):").grid(row=13, column=0, sticky="w", padx=5, pady=5)
    tk.Entry(frame, textvariable=delay_from_var, width=10).grid(row=13, column=1, sticky="w", padx=5, pady=5)

    tk.Label(frame, text="до:").grid(row=13, column=1, sticky="w", padx=95, pady=5)
    tk.Entry(frame, textvariable=delay_to_var, width=10).grid(row=13, column=2, sticky="w", padx=5, pady=5)
    # === PATCH_RANDOM_DELAY_END ===

    send_button = tk.Button(frame, text="НАЧАТЬ РАССЫЛКУ", command=gui_run_sending, width=40)

    send_button.grid(row=14, column=0, columnspan=3, pady=10)

    tk.Button(frame, text="Сохранить настройки", command=save_settings, width=20).grid(
        row=15, column=0, columnspan=3, pady=5
    )

    tk.Label(frame, text="Процесс отправки:").grid(row=16, column=0, sticky="nw", padx=5, pady=5)
    log_text = tk.Text(frame, width=60, height=10, state="disabled")
    # === PATCH_RESIZE_START ===
    log_text.grid(row=16, column=1, columnspan=2, padx=5, pady=5, sticky="nsew")
    # === PATCH_RESIZE_END ===



    load_settings()

    # === PATCH_RESIZE_START ===
    root.resizable(True, True)
    # === PATCH_RESIZE_END ===
    root.mainloop()


if __name__ == "__main__":
    main()
